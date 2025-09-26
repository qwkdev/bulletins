import re

import math
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

GLOBAL_PATH: str = '/home/bulletins/mysite/'
GLOBAL_FONT: str = 'Calibri'

fancyq = {
    "'": ('\u2018', '\u2019'),
    '"': ('\u201C', '\u201D')
}

def format_linebreaks(text: str) -> str:
    return re.sub(r'(?<=<br>)(?=<br>)', '\u2800', text)

def set_table_borders(table, color='000000', size=4, outer=True):
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in (*(['top', 'left', 'bottom', 'right'] if outer else []), 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    table._element.tblPr.append(tblBorders)

def remove_cell_borders(cell):
    tcPr = cell._element.tcPr
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._element.tcPr
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for size, side in zip([top, start, bottom, end], ['top', 'start', 'bottom', 'end']):
        el = tcMar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tcMar.append(el)
        el.set(qn('w:w'), str(size))
        el.set(qn('w:type'), 'dxa')

def set_cell_background(cell, color: str='000000'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)

def zero_paragraph_spacing(cell):
    for p in cell.paragraphs:
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1

def normalize_cell(cell, margins=True, paragraph_spacing=True):
    '''Warning: Doesn't work on empty cells.'''
    cell._tc.remove(cell.paragraphs[0]._p)
    if margins: set_cell_margins(cell, 0, 0, 0, 0)
    if paragraph_spacing: zero_paragraph_spacing(cell)

def normalize_p(p, size, spacing, top=0, bottom=0):
    p.paragraph_format.line_spacing = spacing
    p.paragraph_format.space_before = Pt(top)
    p.paragraph_format.space_after = Pt(bottom)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.name = GLOBAL_FONT

def p_has_image(p):
    for run in p.runs:
        if run._element.xpath('.//w:drawing | .//w:pict'):
            return True
    return False

def remove_blank_p(cell):
    paragraphs = list(cell.paragraphs)
    for p in paragraphs:
        if not p.text.strip() and not p_has_image(p):
            if len(cell.paragraphs) > 1:
                p._element.getparent().remove(p._element)

def safe_normalize_page(left: bool, check: list[bool, bool], table) -> None:
    if not check[0 if left else 1]:
        normalize_cell(table.cell(0, 0 if left else 2))
        check[0 if left else 1] = True

def add_run(p, txt, size, ctx):
    run = p.add_run(txt)
    run.font.size = Pt(size)
    run.font.name = GLOBAL_FONT
    for tag in ctx:
        if tag == 'b':
            run.bold = True
        elif tag == 'i':
            run.italic = True
        elif tag == 'u':
            run.underline = True
        elif tag == 's':
            run.font.superscript = True

def parseText(obj, raw_text, size, spacing, ptop=0, pbottom=0, center=False, left_right=None):
    valid_tags = [
        'br', '_tab',
        'b', 'i', 'u', 's', 'ul',
        '/b', '/i', '/u', '/s', '/ul'
    ]

    text = format_linebreaks(raw_text.replace('<br><ul>', '<ul>').replace('</ul><br>', '</ul>'))
    p = obj.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if left_right is not None:
        p.paragraph_format.tab_stops.add_tab_stop(
            left_right,
            WD_ALIGN_PARAGRAPH.RIGHT
        )
    normalize_p(p, size, spacing, ptop, 0)

    ctx, intag, txt = [], None, ''
    for c in text:
        if c == '<':
            if intag:
                txt += '<'+intag
            intag = ''
        elif c == '>' and intag:
            if txt:
                add_run(p, txt, size, ctx)
                txt = ''

            if intag not in valid_tags:
                add_run(p, '<'+intag+'>', size, ctx)
            elif 'ul' in intag or (ctx and ctx[-1] == 'ul'):
                if intag == '/ul':
                    ctx = ctx[:-1]
                    p = obj.add_paragraph()
                    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    normalize_p(p, size, spacing)
                elif intag == 'ul' or intag == 'br':
                    if intag == 'ul':
                        ctx.append('ul')
                    p = obj.add_paragraph(style="List Bullet")
                    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    normalize_p(p, size, spacing)
                    fmt = p.paragraph_format
                    n = 16
                    m = 1.5

                    fmt.left_indent = Pt(n*m)
                    fmt.first_line_indent = -Pt(n)
            else:
                if intag.startswith('/') and intag[1:] in ctx:
                    ctx = ctx[::-1]
                    ctx.remove(intag[1:])
                    ctx = ctx[::-1]
                elif intag == 'br':
                    p = obj.add_paragraph()
                    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    normalize_p(p, size, spacing)
                elif left_right is not None and intag == '_tab':
                    p.add_run('\t')
                else:
                    ctx.append(intag)

            intag = None
        else:
            if intag is not None:
                intag += c
            else:
                txt += c

    add_run(p, txt, size, ctx)
    if intag: add_run(p, '<'+intag, size, ctx)

    p.paragraph_format.space_after = Pt(pbottom)
    remove_blank_p(obj)

def get_row_height(row):
    row_height = row._tr.trPr.trHeight
    return (row_height.val, row_height.hRule) if row_height is not None else None

def tomm(val: int | float) -> int | float:
    return val / 36000
def topt(val: int | float) -> int | float:
    return val / 12700
def totwips(val: int | float) -> int | float:
    return val / 635
def cellMargin(val: int | float) -> int | float:
    return 350 * val
def toCellMargin(val: int | float) -> int | float:
    return val / 350

def build(
    OUTPUT_PATH: str,
    front_page_margins: tuple[int | float, int | float],
    info_data: list[tuple[int, str]],
    info_size: int | float,
    title: str,
    title_size: int | float,
    church_title: str,
    church_title_size: int | float,
    church_info: str,
    church_info_size: int | float,
    mass_info: list[str],
    mass_info_size: int | float,
    data: list[tuple[int | float, int | float, str]],
    readings: dict,
    reading_margins: tuple[int | float, int | float],
    reading_heading_spacing: int | float,
    reading_heading_size: int | float,
    copyright_size: int | float,
    copyright_spacing: int | float,
    copyright_page: int,
    dpa_page: int
):
    doc = Document()

    style = doc.styles['Normal']
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    margins = Mm(front_page_margins[0]), Mm(front_page_margins[1])

    section.top_margin = round(margins[0] * 0.8)
    section.bottom_margin = round(margins[0] * 0.8)
    section.left_margin = round(margins[1] * 0.8)
    section.right_margin = round(margins[1] * 1.3)
    middle_margin = (margins[1] * .9, margins[1] * .4)

    left_half_width = round(
        (section.page_width / 2)
        - section.left_margin
        - (middle_margin[0] / 2)
    )
    right_half_width = round(
        (section.page_width / 2)
        - section.right_margin
        - (middle_margin[1] / 2)
    )
    a5table = doc.add_table(rows=1, cols=3)
    a5table.autofit = False
    a5table.allow_autofit = False
    a5table.rows[0].height = section.page_height - section.top_margin - section.bottom_margin
    a5table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    a5table.cell(0, 0).width = left_half_width
    a5table.cell(0, 1).width = sum(middle_margin)
    a5table.cell(0, 2).width = right_half_width

    remove_cell_borders(a5table.cell(0, 1))
    normalize_cell(a5table.cell(0, 0))
    normalize_cell(a5table.cell(0, 2))

    set_table_borders(a5table, color='000000', size=4)

    info_data = [(i[0], i[1], i[2].replace('\n', '')) for i in info_data]

    info_table = a5table.cell(0, 2).add_table(rows=len(info_data) + 1, cols=1)
    info_table.autofit = True
    info_table.allow_autofit = True

    total = 0
    for n, ((align, lines, txt), row) in enumerate(zip(info_data, info_table.rows[1:])):
        normalize_cell(row.cells[0], margins=False)

        cell_margin = 70
        set_cell_margins(row.cells[0], cell_margin, 80, cell_margin, 80)
        height = Pt(info_size * 1.22 * lines) + cellMargin(2 * cell_margin)
        if n != len(info_data) - 1:
            row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.cells[0].width = right_half_width
        parseText(row.cells[0], txt, info_size, 1, center=align == 1)

        total += height + cellMargin(2 * cell_margin)

    info_table.rows[0].height = a5table.rows[0].height - total
    front_page = info_table.rows[0].cells[0]

    parseText(front_page, title.replace('\n', ''), title_size, 1.3, 20, center=True)

    logop = front_page.add_paragraph()
    logop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    normalize_p(logop, 1, 1, 5, 0)
    logop.add_run().add_picture(GLOBAL_PATH+'logo.png', width=Mm(54))

    parseText(front_page, church_title.replace('\n', ''), church_title_size, 1.2, 13, center=True)
    parseText(front_page, church_info.replace('\n', ''), church_info_size, 1.2, 2, center=True)

    mass_info = [i.replace('\n', '') for i in mass_info]

    if mass_info:
        match len(mass_info):
            case 1:
                mass_table = front_page.add_table(rows=1, cols=1)
                mass_table_cells = [mass_table.cell(0, 0)]
            case 2:
                mass_table = front_page.add_table(rows=1, cols=2)
                mass_table_cells = [mass_table.cell(0, 0), mass_table.cell(0, 1)]
            case 3:
                mass_table = front_page.add_table(rows=2, cols=2)
                mass_table_cells = [mass_table.cell(0, 0), mass_table.cell(1, 0), mass_table.cell(0, 1).merge(mass_table.cell(1, 1))]
            case 4:
                mass_table = front_page.add_table(rows=2, cols=2)
                mass_table_cells = [mass_table.cell(0, 0), mass_table.cell(1, 0), mass_table.cell(0, 1), mass_table.cell(1, 1)]
            case _:
                cols = math.ceil(len(mass_info) / 2)
                mass_table = front_page.add_table(rows=2, cols=cols)
                mass_table_cells = [mass_table.cell(i, n) for n in range(cols) for i in (0, 1)]

        for cell, txt in zip(mass_table_cells, mass_info):
            normalize_cell(cell, margins=False)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            set_cell_margins(cell, 150, 0, 50, 0)
            parseText(cell, txt, mass_info_size, 1, center=True)

    set_table_borders(info_table, color='000000', size=4, outer=False)

    data = [(i[0], i[1], i[2].replace('\n', '')) for i in data]

    data_table = a5table.cell(0, 0).add_table(rows=len(data), cols=1)
    data_table.autofit = True
    data_table.allow_autofit = True

    for (size, vmargin, txt), row in zip(data, data_table.rows):
        normalize_cell(row.cells[0], margins=False)
        margin = toCellMargin(Mm(vmargin))
        set_cell_margins(row.cells[0], margin, 80, margin, 80)
        row.cells[0].width = left_half_width
        parseText(row.cells[0], txt, size, 1)

    set_table_borders(data_table, color='000000', size=4, outer=False)

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)

    reading_top_margin = Mm(reading_margins[0])
    reading_margin = Mm(reading_margins[1])

    section2.top_margin = reading_top_margin
    section2.bottom_margin = 0
    section2.left_margin = round(reading_margin * 0.8)
    section2.right_margin = round(reading_margin * 1.3)
    middle_margin = (reading_margin * .9, reading_margin * .4)

    left_half_width = round(
        (section2.page_width / 2)
        - section2.left_margin
        - (middle_margin[0] / 2)
    )
    right_half_width = round(
        (section2.page_width / 2)
        - section2.right_margin
        - (middle_margin[1] / 2)
    )

    reading_table = doc.add_table(rows=1, cols=3)
    reading_table.autofit = False
    reading_table.allow_autofit = False
    reading_table.rows[0].height = section2.page_height - reading_top_margin - Mm(8)
    reading_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    reading_table.cell(0, 0).width = left_half_width
    reading_table.cell(0, 1).width = sum(middle_margin)
    reading_table.cell(0, 2).width = right_half_width

    remove_cell_borders(reading_table.cell(0, 1))

    reading_types = {
        'reading1': 'FIRST READING',
        'psalm': 'RESPONSORIAL PSALM',
        'reading2': 'SECOND READING',
        'acclamation': 'GOSPEL ACCLAMATION',
        'gospel': 'GOSPEL'
    }

    reading_page_normalized = [False, False]
    shown_types = []
    for reading in readings:
        if not reading.get('include', True):
            continue
        alt_reading = reading['alt'] and reading['type'] in shown_types
        safe_normalize_page(reading['left'], reading_page_normalized, reading_table)

        reading_page = reading_table.cell(0, 0 if reading['left'] else 2)
        parseText(reading_page, '<b>'
            + ('OR' if alt_reading else reading_types[reading['type']])
            + ('</b>  <i>wording may differ if sung</i>' if reading['type'] in ['psalm', 'acclamation'] and not alt_reading and reading['sameline'] else '</b>')
            + '<_tab>'
            + reading['ref'],
        reading_heading_size, 1, pbottom=reading_heading_spacing, left_right=left_half_width)
        shown_types.append(reading['type'])

        if reading['title']:
            parseText(reading_page, '<b><i>' + fancyq['"'][0] + reading['title'] + fancyq['"'][1] + '</i></b>', reading_heading_size, 1, pbottom=reading_heading_spacing)
        if reading['type'] in ['reading1', 'reading2', 'gospel']:
            parseText(reading_page, reading['text'], reading['size'], 1, pbottom=reading['margin'])
        if reading['type'] in ['psalm', 'acclamation']:
            if not reading['sameline'] and not alt_reading:
                parseText(reading_page, '<i>wording may differ if sung</i>', reading_heading_size, 1, pbottom=reading_heading_spacing)

            if reading['type'] == 'psalm':
                parseText(reading_page, '<b>' + reading['text'][0] + '</b>', reading['size'], 1, pbottom=reading_heading_spacing)
                parseText(reading_page, reading['text'][1], reading['size'], 1, pbottom=reading['margin'])
            else:
                parseText(reading_page, '<b>Alleluia, alleluia.</b><br>' + reading['text'] + '<br><b>Alleluia.</b>', reading['size'], 1, pbottom=reading['margin'])

    safe_normalize_page(copyright_page == 0, reading_page_normalized, reading_table)
    safe_normalize_page(dpa_page == 0, reading_page_normalized, reading_table)

    parseText(reading_table.cell(0, 0 if copyright_page == 0 else 2),
        '''<i>The text of Sacred Scripture in the Lectionary is from the English Standard Version of the Bible, Catholic Edition (ESV-CE), published by Asian Trading Corporation, \u00a9 2017 Crossway. All rights are reserved. The English Standard Version of the Bible, Catholic Edition is published in the United Kingdom by SPCK Publishing. The Psalms and Canticles are from Abbey Psalms and Canticles \u00a9 2018 United States Conference of Catholic Bishops. Reprinted with permission.</i>''',
        copyright_size, 1, pbottom=copyright_spacing)
    parseText(reading_table.cell(0, 0 if dpa_page == 0 else 2),
        '''<i>Please note the Data Protection Act 2018 restricts the inclusion of the names of our sick unless their consent is given. If you wish to include someone\u2019s name here please speak to Fr John on completing a Consent Form from the sacristy.</i>''',
        copyright_size, 1)

    doc.save(GLOBAL_PATH+OUTPUT_PATH)
